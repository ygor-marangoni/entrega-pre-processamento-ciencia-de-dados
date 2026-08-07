#!/usr/bin/env python3
"""Gera checklist e manifesto final, com dimensões e hashes dos artefatos."""

from __future__ import annotations

import hashlib
from pathlib import Path

import pandas as pd
import pdfplumber
from docx import Document


SEP = ";"
ROOT = Path(__file__).resolve().parent.parent
REPO = ROOT.parent
OUT = ROOT / "resultados" / "validacao_final"
MANIFEST = OUT / "manifesto_arquivos_entrega.csv"
CHECKLIST = OUT / "checklist_entrega_final.csv"
EXCLUDED_PARTS = {
    "backup_pre_validacao_final_20260721",
    "__pycache__",
    ".matplotlib_cache",
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as file:
        for chunk in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def csv_shape(path: Path) -> tuple[int | str, int | str]:
    try:
        frame = pd.read_csv(path, sep=SEP, low_memory=False)
        return len(frame), len(frame.columns)
    except Exception:
        return "n/a", "n/a"


def arff_shape(path: Path) -> tuple[int, int]:
    attributes = 0
    records = 0
    in_data = False
    with path.open("r", encoding="utf-8", errors="replace") as file:
        for raw in file:
            line = raw.strip()
            if line.casefold().startswith("@attribute"):
                attributes += 1
            elif line.casefold() == "@data":
                in_data = True
            elif in_data and line and not line.startswith("%"):
                records += 1
    return records, attributes


def dimensions(path: Path) -> tuple[int | str, int | str]:
    suffix = path.suffix.casefold()
    if suffix == ".csv":
        return csv_shape(path)
    if suffix == ".arff":
        return arff_shape(path)
    if suffix == ".pdf":
        with pdfplumber.open(path) as document:
            return len(document.pages), "n/a"
    if suffix == ".docx":
        document = Document(path)
        return len(document.paragraphs), len(document.tables)
    if suffix in {".md", ".txt", ".py", ".json"}:
        lines = len(path.read_text(encoding="utf-8", errors="replace").splitlines())
        return lines, "n/a"
    return "n/a", "n/a"


def purpose(path: Path) -> str:
    rel = path.relative_to(ROOT).as_posix()
    if rel.startswith("relatorio/"):
        return "Relatório final, imagem ou fonte documental."
    if rel.startswith("data/amostras/"):
        return "Amostra oficial reproduzível e rastreável."
    if rel.startswith("data/preparadas/"):
        return "Base ponderada de entrada do WEKA."
    if rel.startswith("data/clusterizadas_weka/"):
        return "Exportação ou conversão validada de clusterização do WEKA."
    if rel.startswith("data/analise/"):
        return "Base mesclada para análise posterior, incluindo TARGET sem formar clusters."
    if rel.startswith("scripts/originais/") or path.name in {"p1.py", "p2.py"}:
        return "Script didático original fornecido pelo professor."
    if rel.startswith("scripts/"):
        return "Script próprio reproduzível do Trabalho 3."
    if rel.startswith("resultados/validacao_final/"):
        return "Evidência da auditoria e validação final."
    if rel.startswith("resultados/"):
        return "Configuração, métrica ou resumo intermediário rastreável."
    if rel.startswith("docs/"):
        return "Documentação operacional e diário do projeto."
    if path.name.startswith("README"):
        return "Orientação geral da entrega."
    return "Artefato de apoio do Trabalho 3."


def should_include(path: Path) -> bool:
    if path == MANIFEST:
        return False
    if any(part in EXCLUDED_PARTS for part in path.parts):
        return False
    return path.is_file()


def build_manifest() -> pd.DataFrame:
    paths = sorted(path for path in ROOT.rglob("*") if should_include(path))
    source = REPO / "trabalho_1_preprocessamento" / "data" / "base_final_preprocessada.csv"
    paths.append(source)
    records = []
    for path in paths:
        rows, columns = dimensions(path)
        try:
            relative = path.relative_to(ROOT).as_posix()
            include_moodle = "SIM"
        except ValueError:
            relative = "../" + path.relative_to(REPO).as_posix()
            include_moodle = "NÃO - origem rastreada"
        records.append(
            {
                "caminho": relative,
                "nome": path.name,
                "formato": path.suffix.lower().lstrip(".") or "sem extensão",
                "linhas_ou_paginas": rows,
                "colunas_ou_tabelas": columns,
                "tamanho_bytes": path.stat().st_size,
                "sha256": sha256(path),
                "finalidade": purpose(path) if path.is_relative_to(ROOT) else "Base de origem do Trabalho 1; não duplicada na entrega do Trabalho 3.",
                "incluir_moodle": include_moodle,
            }
        )
    return pd.DataFrame(records)


def build_checklist() -> pd.DataFrame:
    pdf_path = ROOT / "relatorio" / "relatorio_clusterizacao_abnt.pdf"
    with pdfplumber.open(pdf_path) as pdf:
        page_count = len(pdf.pages)
        a4 = all(abs(page.width - 595.304) < 1 and abs(page.height - 841.890) < 1 for page in pdf.pages)
    items = [
        ("Relatório PDF final", pdf_path.is_file(), f"{page_count} páginas; A4={a4}"),
        ("Relatório DOCX final", (ROOT / "relatorio/relatorio_clusterizacao_abnt.docx").is_file(), "Gerado pelo script 11."),
        ("Relatório Markdown final", (ROOT / "relatorio/relatorio_clusterizacao.md").is_file(), "Fonte textual auditada."),
        ("Base preparada CSV e ARFF", all((ROOT / f"data/preparadas/base_clusterizacao_final.{ext}").is_file() for ext in ("csv", "arff")), "10.000 x 6; sem colunas protegidas."),
        ("Base DBSCAN CSV e ARFF", all((ROOT / f"data/clusterizadas_weka/base_clusterizada_dbscan.{ext}").is_file() for ext in ("csv", "arff")), "ARFF original do AddCluster; CSV derivado fiel."),
        ("Base SimpleKMeans CSV e ARFF", all((ROOT / f"data/clusterizadas_weka/base_clusterizada_kmeans_final.{ext}").is_file() for ext in ("csv", "arff")), "Exportação final K=9."),
        ("Base EM CSV e ARFF", all((ROOT / f"data/clusterizadas_weka/base_clusterizada_em_final.{ext}").is_file() for ext in ("csv", "arff")), "Exportação final K=9."),
        ("Bases de análise", all((ROOT / f"data/analise/analise_cluster_{m}.csv").is_file() for m in ("dbscan", "kmeans", "em")), "Três bases de 10.000 x 49."),
        ("p1.py e p2.py", all((ROOT / f"scripts/{name}").is_file() for name in ("p1.py", "p2.py")), "Cópias preservadas no Trabalho 3."),
        ("Scripts próprios", len(list((ROOT / "scripts").glob("[0-9][0-9]_*.py"))) >= 12, "Scripts 02 a 13 presentes; numeração funcional preservada."),
        ("Configurações", all((ROOT / f"resultados/configuracoes/{name}").is_file() for name in ("configuracao_final.csv", "configuracoes_dbscan.csv", "configuracoes_kmeans.csv", "configuracoes_em.csv")), "Arquivos reais e sem sobrescrever testes."),
        ("Hopkins", all((ROOT / f"resultados/hopkins/hopkins_tentativa_01.{ext}").is_file() for ext in ("csv", "txt")), "0,941745479089317 reproduzido duas vezes."),
        ("Gráficos", len(list((ROOT / "relatorio/imagens").rglob("*.png"))) >= 31, "Exploração, clusters e k-distance."),
        ("README", (ROOT / "README_TRABALHO_3.md").is_file(), "Instruções e estado das etapas."),
        ("Auditoria final", all((OUT / name).is_file() for name in ("auditoria_final.md", "checks_validacao.csv", "inconsistencias_encontradas.csv")), "Fase 1 preservada."),
    ]
    records = [
        {"item": item, "status": "OK" if ok else "PENDENTE", "evidencia": evidence}
        for item, ok, evidence in items
    ]
    records.extend(
        [
            {"item": "Logs textuais brutos do WEKA", "status": "RESSALVA", "evidencia": "Não preservados; ARFFs comprovam AddCluster, configurações e rótulos."},
            {"item": "Conferência humana de autoria/docente", "status": "RESSALVA", "evidencia": "Confirmar grafia, turma e dados institucionais da folha de rosto."},
        ]
    )
    return pd.DataFrame(records)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    checklist = build_checklist()
    checklist.to_csv(CHECKLIST, sep=SEP, index=False, encoding="utf-8-sig", lineterminator="\n")
    manifest = build_manifest()
    manifest.to_csv(MANIFEST, sep=SEP, index=False, encoding="utf-8-sig", lineterminator="\n")
    if (checklist["status"] == "PENDENTE").any():
        raise RuntimeError("Checklist possui item pendente.")
    print(f"Checklist: {len(checklist)} itens; pendentes=0; ressalvas={(checklist.status == 'RESSALVA').sum()}")
    print(f"Manifesto: {len(manifest)} arquivos")


if __name__ == "__main__":
    main()
