#!/usr/bin/env python3
"""Converte o relatório Markdown validado em DOCX acadêmico no padrão ABNT."""

from __future__ import annotations

import argparse
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SCRIPT_DIR = Path(__file__).resolve().parent
ROOT = SCRIPT_DIR.parent
DEFAULT_MARKDOWN = ROOT / "relatorio" / "relatorio_clusterizacao.md"
DEFAULT_OUTPUT = ROOT / "relatorio" / "relatorio_clusterizacao_abnt.docx"
FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
CONTENT_DXA = 9072  # 16 cm: A4 menos margens esquerda/direita de 3/2 cm.
TABLE_CAPTIONS = [
    "Características e dimensões das bases utilizadas",
    "Atributos selecionados, tipos e pesos",
    "Resultado da Estatística de Hopkins",
    "Configurações testadas no DBSCAN",
    "Configurações testadas no SimpleKMeans",
    "Configurações testadas no EM",
    "Comparação das métricas internas",
    "Perfis recorrentes identificados",
]
TOC_ENTRIES = [
    ("1 Introdução", 1),
    ("2 Descrição do problema", 2),
    ("3 Descrição da base", 3),
    ("4 Criação da amostra", 4),
    ("5 Seleção dos atributos", 5),
    ("6 Análise da dispersão", 6),
    ("7 Pesos e transformações", 9),
    ("8 Estatística de Hopkins", 10),
    ("9 Descrição do DBSCAN", 11),
    ("10 Descrição do SimpleKMeans", 12),
    ("11 Descrição do EM", 13),
    ("12 Configurações utilizadas", 14),
    ("13 Resultados do DBSCAN", 16),
    ("14 Resultados do K-Means", 18),
    ("15 Resultados do EM", 20),
    ("16 Comparação dos métodos", 22),
    ("17 Perfis identificados", 24),
    ("18 Aplicações comerciais", 25),
    ("19 Limitações", 26),
    ("20 Conclusão", 27),
    ("21 Referências", 28),
    ("22 Apêndice A — Arquivos e rastreabilidade", 30),
    ("23 Apêndice B — Configurações completas", 31),
    ("24 Apêndice C — Bases anexadas", 32),
]

# Seções que ganharam uma página de texto após a validação final. O
# incremento ocorre antes do próximo título principal para manter paginação
# contínua e sumário estático coerente no PDF exportado pelo LibreOffice.
EXTRA_BODY_PAGES_BEFORE_HEADING = {13: 1, 17: 1, 22: 1}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--markdown", type=Path, default=DEFAULT_MARKDOWN)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--sobrescrever", action="store_true")
    return parser.parse_args()


def set_run_font(run, size: float = 12, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"Larguras da tabela não somam {CONTENT_DXA}: {widths}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "100")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Cm(widths[index] / 1440 * 2.54)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def proportional_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    maxima = [max(5, min(45, max(len(row[index]) for row in rows))) for index in range(cols)]
    if cols >= 5:
        maxima = [min(value, 20) for value in maxima]
    total = sum(maxima)
    widths = [max(700, round(CONTENT_DXA * value / total)) for value in maxima]
    difference = CONTENT_DXA - sum(widths)
    widths[-1] += difference
    if widths[-1] < 700:
        deficit = 700 - widths[-1]
        widths[-1] = 700
        largest = max(range(cols - 1), key=lambda index: widths[index])
        widths[largest] -= deficit
    return widths


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE \\* Arabic")
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    run_properties.extend((fonts, size))
    text = OxmlElement("w:t")
    text.text = "1"
    run.extend((run_properties, text))
    field.append(run)
    paragraph._p.append(field)


def add_static_page_number(paragraph, value: int) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(paragraph.add_run(str(value)), 10)


def set_page_number_start(section, value: int) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(value))


def configure_section(section, body: bool = False, page_number: int | None = None) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header.paragraphs[0].clear()
    section.footer.paragraphs[0].clear()
    if body:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run("Trabalho Prático 3 - Clusterização")
        set_run_font(run, 9, italic=True)
        if page_number is None:
            add_page_number(section.footer.paragraphs[0])
        else:
            # Use a real PAGE field even when a section restarts the counter.
            # Static two-digit text was clipped by LibreOffice on alternating
            # sections (for example, 11 appeared visually as 1).
            add_page_number(section.footer.paragraphs[0])
            set_page_number_start(section, page_number)


def define_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5

    tokens = {
        "Heading 1": (14, 18, 10, True),
        "Heading 2": (12, 12, 6, True),
        "Heading 3": (12, 8, 4, True),
    }
    for name, (size, before, after, page_break) in tokens.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True

    for style_name in ("Caption", "Intense Quote"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.color.rgb = BLACK

    caption = doc.styles["Caption"]
    caption.font.size = Pt(10)
    caption.font.italic = False
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.line_spacing = 1.0

    if "Fonte de tabela/figura" not in doc.styles:
        source = doc.styles.add_style("Fonte de tabela/figura", WD_STYLE_TYPE.PARAGRAPH)
    else:
        source = doc.styles["Fonte de tabela/figura"]
    source.font.name = FONT
    source._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    source._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    source.font.size = Pt(10)
    source.font.italic = True
    source.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    source.paragraph_format.first_line_indent = Cm(0)
    source.paragraph_format.space_before = Pt(2)
    source.paragraph_format.space_after = Pt(12)
    source.paragraph_format.line_spacing = 1.0

    if "Sumário" not in doc.styles:
        toc_heading = doc.styles.add_style("Sumário", WD_STYLE_TYPE.PARAGRAPH)
    else:
        toc_heading = doc.styles["Sumário"]
    toc_heading.font.name = FONT
    toc_heading._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    toc_heading._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    toc_heading.font.size = Pt(14)
    toc_heading.font.bold = True
    toc_heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_heading.paragraph_format.space_after = Pt(18)


def add_numbering(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    level.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    paragraph.paragraph_format.first_line_indent = Cm(0)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_element)
    p_pr.append(num_pr)


def add_inline(paragraph, text: str, size: float = 12) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            set_run_font(paragraph.add_run(text[position:match.start()]), size)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size - 0.5)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size, italic=True)
        position = match.end()
    if position < len(text):
        set_run_font(paragraph.add_run(text[position:]), size)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section)
    for text, size, bold, after in [
        ("UNIVERSIDADE FEDERAL DE UBERLÂNDIA", 12, True, 0),
        ("FACULDADE DE COMPUTAÇÃO", 12, True, 54),
        ("Gil Antony Borba\nRaphael Muniz Varela\nVictor Leal\nYgor Marangoni", 12, False, 90),
        ("RELATÓRIO DE CLUSTERIZAÇÃO PARA ANÁLISE DE PERFIS DE CRÉDITO", 14, True, 110),
        ("Monte Carmelo - MG\n2026", 12, False, 0),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        set_run_font(run, size, bold=bold)


def add_title_page(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(82)
    run = p.add_run("Gil Antony Borba\nRaphael Muniz Varela\nVictor Leal\nYgor Marangoni")
    set_run_font(run, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(72)
    run = p.add_run("RELATÓRIO DE CLUSTERIZAÇÃO PARA ANÁLISE DE PERFIS DE CRÉDITO")
    set_run_font(run, 14, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(8)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(18)
    add_inline(p, "Trabalho Prático 3 apresentado à disciplina de Ciência de Dados da Universidade Federal de Uberlândia, como requisito parcial de avaliação.", 10)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(8)
    p.paragraph_format.first_line_indent = Cm(0)
    add_inline(p, "Professor: Carlos Cesar Mansur Tuma", 10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(110)
    run = p.add_run("Monte Carmelo - MG\n2026")
    set_run_font(run, 12)


def add_toc(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section)
    doc.add_paragraph("SUMÁRIO", style="Sumário")
    for title, page in TOC_ENTRIES:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(15.7), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        set_run_font(p.add_run(f"{title}\t{page}"), 9)


def add_table(doc: Document, rows: list[list[str]], table_number: int) -> None:
    caption = doc.add_paragraph(style="Caption")
    caption.paragraph_format.keep_with_next = True
    add_inline(caption, f"Tabela {table_number} - {TABLE_CAPTIONS[table_number - 1]}", 10)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = proportional_widths(rows)
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, value, 9 if len(rows[0]) >= 6 else 10)
            if row_index == 0:
                shade_cell(cell, "E7E6E6")
                for run in p.runs:
                    run.bold = True
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, widths)


def parse_markdown_body(doc: Document, markdown: Path) -> None:
    text = markdown.read_text(encoding="utf-8")
    separators = [match.end() for match in re.finditer(r"(?m)^---\s*$", text)]
    if len(separators) < 2:
        raise ValueError("Markdown não possui os separadores esperados da capa e do sumário.")
    lines = text[separators[1]:].strip().splitlines()
    bullet_num_id = add_numbering(doc, "bullet")
    decimal_num_id = add_numbering(doc, "decimal")
    table_number = 0
    image_number = 0
    heading_one_number = 0
    body_page = 1
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        image = re.match(r"^!\[(.+)]\((.+)\)$", line)
        if heading:
            level = len(heading.group(1))
            if level == 1:
                heading_one_number += 1
                if heading_one_number > 1:
                    body_page += 1 + EXTRA_BODY_PAGES_BEFORE_HEADING.get(heading_one_number, 0)
                    configure_section(
                        doc.add_section(WD_SECTION.NEW_PAGE),
                        body=True,
                        page_number=body_page,
                    )
            doc.add_paragraph(heading.group(2), style=f"Heading {level}")
            index += 1
            continue
        if image:
            image_number += 1
            if image_number in {2, 4, 6, 8, 10}:
                body_page += 1
                configure_section(
                    doc.add_section(WD_SECTION.NEW_PAGE),
                    body=True,
                    page_number=body_page,
                )
            image_path = (markdown.parent / image.group(2)).resolve()
            if not image_path.exists():
                raise FileNotFoundError(image_path)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            picture = p.add_run().add_picture(str(image_path), width=Cm(15.2))
            picture._inline.docPr.set("descr", image.group(1))
            picture._inline.docPr.set("title", f"Figura {image_number}")
            caption = doc.add_paragraph(style="Caption")
            caption.paragraph_format.keep_with_next = True
            add_inline(caption, image.group(1), 10)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[index + 1]):
            rows = [[cell.strip() for cell in line.strip("|").split("|")]]
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                rows.append([cell.strip() for cell in lines[index].strip("|").split("|")])
                index += 1
            if any(len(row) != len(rows[0]) for row in rows):
                raise ValueError(f"Tabela Markdown irregular próxima de: {rows[0]}")
            table_number += 1
            add_table(doc, rows, table_number)
            continue
        bullet = re.match(r"^-\s+(.+)$", line)
        decimal = re.match(r"^\d+\.\s+(.+)$", line)
        if bullet or decimal:
            p = doc.add_paragraph()
            apply_num(p, bullet_num_id if bullet else decimal_num_id)
            add_inline(p, (bullet or decimal).group(1), 12)
            index += 1
            continue
        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            candidate = lines[index].rstrip()
            if not candidate:
                index += 1
                break
            if re.match(r"^(#{1,3})\s+", candidate) or candidate.startswith("|") or candidate.startswith("- ") or re.match(r"^\d+\.\s+", candidate) or candidate.startswith("!["):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph_text = " ".join(paragraph_lines)
        style = "Fonte de tabela/figura" if paragraph_text.startswith("Fonte:") else None
        p = doc.add_paragraph(style=style)
        code_tokens = re.findall(r"`[^`]+`", paragraph_text)
        long_technical_token = any(
            len(token) > 35 and ("_" in token or "/" in token)
            for token in paragraph_text.split()
        )
        if style is None and (
            re.search(r"\b[a-f0-9]{40,}\b", paragraph_text, flags=re.IGNORECASE)
            or len(code_tokens) >= 2
            or long_technical_token
        ):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline(p, paragraph_text, 10 if style else 12)
    if table_number != len(TABLE_CAPTIONS):
        raise ValueError(f"Esperadas {len(TABLE_CAPTIONS)} tabelas; encontradas {table_number}.")
    if heading_one_number != 24 or image_number != 12 or body_page != 32:
        raise ValueError(
            "Paginação inesperada: "
            f"{heading_one_number} seções, {image_number} figuras e página final {body_page}."
        )


def enable_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def validate_document(path: Path) -> None:
    doc = Document(path)
    heading_ones = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
    if len(heading_ones) != 24:
        raise ValueError(f"Esperadas 24 seções principais; encontradas {len(heading_ones)}.")
    if len(doc.tables) != 8:
        raise ValueError(f"Esperadas 8 tabelas; encontradas {len(doc.tables)}.")
    if len(doc.inline_shapes) != 12:
        raise ValueError(f"Esperadas 12 figuras; encontradas {len(doc.inline_shapes)}.")
    table_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    all_text = "\n".join(p.text for p in doc.paragraphs) + "\n" + table_text
    for forbidden in ("[[TOC]]", "PREENCHER", "Lorem ipsum", "Atualize o sumário"):
        if forbidden in all_text:
            raise ValueError(f"Placeholder proibido encontrado: {forbidden}")
    toc_text = {p.text for p in doc.paragraphs}
    missing_toc = [f"{title}\t{page}" for title, page in TOC_ENTRIES if f"{title}\t{page}" not in toc_text]
    if missing_toc:
        raise ValueError(f"Entradas ausentes no sumário: {missing_toc}")
    if "TARGET" not in all_text or "0,941745479089" not in all_text:
        raise ValueError("Conteúdo obrigatório ausente no DOCX.")
    for section in doc.sections:
        if abs(section.page_width.cm - 21) > 0.02 or abs(section.page_height.cm - 29.7) > 0.02:
            raise ValueError("Seção fora do formato A4.")


def main() -> None:
    args = parse_args()
    markdown = args.markdown.resolve()
    output = args.output.resolve()
    if not markdown.exists():
        raise FileNotFoundError(markdown)
    if output.exists() and not args.sobrescrever:
        raise FileExistsError(f"Saída já existe: {output}. Use --sobrescrever.")
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    define_styles(doc)
    add_cover(doc)
    add_title_page(doc)
    add_toc(doc)
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(body_section, body=True, page_number=1)
    parse_markdown_body(doc, markdown)
    enable_update_fields(doc)
    doc.core_properties.title = "Relatório de Clusterização para Análise de Perfis de Crédito"
    doc.core_properties.subject = "Trabalho Prático 3 de Ciência de Dados"
    doc.core_properties.author = "Gil Antony Borba; Raphael Muniz Varela; Victor Leal; Ygor Marangoni"
    temporary = output.with_suffix(output.suffix + ".tmp")
    try:
        doc.save(temporary)
        os.replace(temporary, output)
    finally:
        temporary.unlink(missing_ok=True)
    validate_document(output)
    print(f"Relatório gerado: {output}")
    print("Validação: 24 seções, 8 tabelas, 12 figuras, paginação até 32, A4 e sem placeholders.")


if __name__ == "__main__":
    main()
