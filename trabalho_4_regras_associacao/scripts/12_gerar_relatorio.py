"""Gera o relatório detalhado do Trabalho 4 em Markdown e DOCX.

O DOCX desta etapa é uma versão de conteúdo revisável. A padronização ABNT e o
PDF são deliberadamente reservados às etapas posteriores.
"""
from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Mm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
REL = ROOT / "relatorio"


def csv_rows(path: Path):
    # Alguns CSVs do PowerShell foram salvos com BOM UTF-8.
    with path.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f, delimiter=";"))


def item(text: str) -> str:
    return text.replace("__", " = ").replace("_", " ")


def table_md(headers, rows):
    out = ["| " + " | ".join(headers) + " |", "|" + "|".join(["---"] * len(headers)) + "|"]
    out += ["| " + " | ".join(str(x).replace("|", "/") for x in row) + " |" for row in rows]
    return "\n".join(out)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    code = OxmlElement('w:instrText'); code.set(qn('xml:space'), 'preserve'); code.text = instruction
    separate = OxmlElement('w:fldChar'); separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t'); text.text = 'Atualize o campo no Word.'
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, code, separate, text, end])


def shade(cell, fill='D9D9D9'):
    props = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); props.append(shd)


def add_doc_table(doc, headers, rows, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = str(text)
        shade(cell)
        for run in cell.paragraphs[0].runs:
            run.bold = True; run.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = str(text)
            for p in cell.paragraphs:
                for run in p.runs: run.font.size = Pt(font_size)
    source = doc.add_paragraph("Fonte: elaboração própria a partir dos arquivos auditáveis do Trabalho 4.")
    source.runs[0].italic = True; source.paragraph_format.space_before = Pt(3); source.paragraph_format.space_after = Pt(6)


def main():
    top = csv_rows(ROOT / "resultados/regras/analise_top20_regras.csv")
    support = csv_rows(ROOT / "resultados/apriori/testes_suporte/resumo_testes_suporte.csv")
    closed = csv_rows(ROOT / "resultados/conjunto_fechado/auditoria_fechamento.csv")
    decisions = csv_rows(ROOT / "resultados/discretizacao/proposta_faixas.csv")
    frequency = csv_rows(ROOT / "resultados/discretizacao/frequencias_finais.csv")
    n_closed = sum(r["fechado"] == "SIM" for r in closed)
    n_not_closed = len(closed) - n_closed
    classes = {c: sum(r["Classificação"] == c for r in top) for c in ("ÓBVIA", "INTERESSANTE", "NOVIDADE")}

    section_text = [
        ("1 Introdução", "Este relatório apresenta a mineração de regras de associação Apriori sobre uma amostra reproduzível de clientes. O propósito é descrever coocorrências entre características, não prever TARGET e não atribuir causalidade. Todas as quantidades apresentadas foram extraídas de arquivos preservados no repositório."),
        ("2 Descrição do problema", "O enunciado exige discretizar atributos, executar Apriori no WEKA, preservar itemsets, construir o conjunto fechado, ordenar regras por Lift e interpretar vinte regras. TARGET, SK_ID_CURR e ROW_ID_AMOSTRA foram excluídos da mineração."),
        ("3 Contexto dos trabalhos anteriores", "O Trabalho 1 produziu a base final pré-processada com 307.511 registros e 41 colunas. O Trabalho 3 gerou a amostra analítica reproduzível de 10.000 registros. A base ponderada de clusterização não foi usada, porque contém transformações e pesos específicos daquele trabalho."),
        ("4 Descrição da base", "A origem aprovada é trabalho_3_clusterizacao/data/amostras/base_amostra_10000_analise.csv. Ela contém valores originais não ponderados e mantém identificadores apenas para rastreabilidade. A representação Apriori final usa 10.000 transações e oito dimensões conceituais."),
        ("5 Escolha da amostra", "A amostra do Trabalho 3 foi escolhida por ser reproduzível, possuir 10.000 linhas e não conter Min-Max, raiz de peso ou peso de clusterização. A comparação de linhas confirmou que ela preserva valores de origem antes da ponderação."),
        ("6 Auditoria dos atributos", "Foram auditados 13 candidatos quanto a tipo, distribuição, quantis, ausências, zeros, assimetria, concentração e outliers. A seleção buscou interpretação, diversidade e frequência suficiente, evitando categorias extremamente raras."),
        ("7 Seleção dos oito atributos", "Foram selecionados crédito, renda, idade, filhos, posse de carro, situação familiar, créditos ativos e taxa de rejeição prévia. As dimensões combinam perfil demográfico, familiar e financeiro, sem incluir o alvo ou identificadores."),
        ("8 Atributos descartados e motivos", "CREDIT_INCOME_RATIO foi descartado por redundância matemática quando crédito e renda já estão presentes. SER_DIVIDA_ATRASADA foi descartado pela concentração próxima de 98,8% em zero. SER_QTDE_EMPRESTIMOS, PREV_QTDE_TENTATIVAS e o candidato remanescente não foram selecionados para preservar diversidade e limitar a oito dimensões."),
        ("9 Análise da distribuição", "As faixas foram definidas por percentis, quartis, semântica e tamanho de categoria. A validação evitou classes abaixo de 5%, reduzindo a chance de regras artificiais sustentadas por poucos registros."),
        ("10 Processo de discretização", "As oito variáveis foram convertidas em categorias interpretáveis. Variáveis monetárias receberam quatro faixas; idade recebeu cinco grupos; filhos, carro, situação familiar, créditos ativos e rejeição receberam categorias semanticamente nomeadas."),
        ("11 Justificativa individual das faixas", "Crédito e renda usam cortes robustos à presença de extremos. Idade usa intervalos interpretáveis. Filhos separa sem filhos, um filho e dois ou mais. Categorias financeiras distinguem ausência, intensidade moderada e intensidade maior sem deixar o zero numérico como valor de negócio."),
        ("12 Problema de treatZeroAsMissing", "A auditoria do WEKA 3.8.7 confirmou que -Z trata o primeiro valor nominal como ausente, e não o texto zero. Portanto, o ARFF nominal multivalorado não foi minerado com -Z, pois isso eliminaria categorias reais."),
        ("13 Tratamento dos zeros semânticos", "Foi criada uma representação one-hot com 27 atributos {0,1}. Em cada coluna, 0 significa item ausente e 1 item presente. Assim, -Z ignora somente ausências. Cada transação possui exatamente oito valores 1, um por dimensão original."),
        ("14 Geração da base final", "A base semântica discretizada foi preservada sem alteração. A base técnica binária tem 10.000 linhas, 27 itens e nenhuma coluna de TARGET, identificador, cluster ou peso. As frequências binárias foram confrontadas com as frequências nominais."),
        ("15 Validação das categorias", "Foram validados 27 itens, 10.000 linhas, valores apenas 0 e 1, oito itens ativos por transação, ordem preservada e igualdade exata das frequências. Nenhuma categoria válida é representada pelo valor 0."),
        ("16 Fundamentos de regras de associação", "Uma regra X -> Y descreve a coocorrência de conjuntos de itens em transações. Ela não estabelece que X cause Y. A interpretação depende simultaneamente de suporte, confiança, Lift, contexto de negócio e limitações amostrais."),
        ("17 Método Apriori", "Apriori explora a propriedade de que um itemset frequente possui subconjuntos frequentes. O WEKA listou os conjuntos frequentes por tamanho e, a partir deles, gerou regras ordenadas pela métrica configurada."),
        ("18 Conceito de suporte", "Suporte é a proporção de transações que contêm todos os itens da regra. Na saída estruturada, o suporte relativo é o suporte absoluto informado pelo WEKA dividido por 10.000."),
        ("19 Conceito de confiança", "Confiança é a proporção das transações com antecedente que também apresentam o consequente. Confiança alta pode refletir uma categoria consequente muito frequente; por isso não deve ser usada isoladamente."),
        ("20 Conceito de Lift", "Lift próximo de 1 sugere independência aproximada; maior que 1 sugere associação positiva e menor que 1, associação negativa. Lift não expressa causalidade e deve ser lido junto ao suporte e à confiança."),
        ("21 Configurações do WEKA", "Foi usado WEKA 3.8.7 com Java 17 e opção de compatibilidade --add-opens java.base/java.lang=ALL-UNNAMED. A execução final usou -N 30, -T 1 (Lift), -C 0.00, -M 0.01, -U 0.01, -D 0.01, -I e -Z."),
        ("22 Busca do lowerBoundMinSupport", "A busca começou com suporte alto e reduziu progressivamente o limite, preservando cada tentativa. Com a confiança padrão de 0,90, nenhuma configuração retornou trinta regras de três itens. Com autorização posterior, foram investigadas outras métricas e pontuações mínimas."),
        ("23 Todas as tentativas de suporte", "O arquivo resumo_testes_suporte.csv preserva cada tentativa, parâmetros, tempo, itemsets e contagem de regras de três itens. A tabela de apêndice reproduz esse inventário, incluindo as execuções exploratórias explicitamente identificadas."),
        ("24 Escolha do suporte final", "O suporte efetivo de 0,01, com Lift e pontuação mínima 0,00, retornou 1.504 regras na exploração N=2000, das quais 36 têm três itens. Essa foi a primeira configuração a disponibilizar ao menos trinta regras válidas de três itens."),
        ("25 Geração das 30 regras", "A execução obrigatória com N=30 foi preservada em resultado_apriori_final.txt. Ela gerou 30 regras e 5.119 itemsets em 4,22 segundos. As trinta regras são de cinco ou seis itens porque o ranking por Lift priorizou regras longas."),
        ("26 Restrição de três itens", "O WEKA não oferece opção nativa para restringir o total de itens da regra. Por transparência, a saída N=30 não foi falsificada. As regras de três itens foram identificadas e estruturadas a partir da execução exploratória, mantida separada e auditável."),
        ("27 Geração dos itemsets", "A execução final listou 27 itemsets L(1), 309 L(2), 1.405 L(3), 2.103 L(4), 1.093 L(5), 176 L(6) e 6 L(7), totalizando 5.119."),
        ("28 Conceito de itemset fechado", "Um itemset X é fechado quando não existe superconjunto próprio Y com o mesmo suporte. O conjunto fechado reduz redundância sem eliminar a informação de suporte associada a extensões idênticas."),
        ("29 Construção do conjunto fechado", f"A auditoria comparou cada itemset com supersets de mesmo suporte absoluto. Dos 5.119 itemsets, {n_closed} foram fechados e {n_not_closed} foram removidos como não fechados, sempre com superconjunto-testemunha registrado."),
        ("30 Regras removidas por redundância", "Todas as 30 regras finais correspondem a itemsets fechados. Entre as 1.504 regras exploratórias, 1.412 permaneceram; as 36 regras de três itens permanecem no conjunto fechado."),
        ("31 Ordenação final por Lift", "A seleção considerou regras exploratórias fechadas com exatamente três itens e métricas finitas. O ranking foi Lift decrescente, seguido de suporte e confiança decrescentes. Havia 36 candidatas e foram selecionadas as vinte primeiras."),
        ("32 Top 20 regras", "O Top 20 tem suporte de 1% em todas as regras e Lift entre 1,11 e 0,95. A manutenção de Lifts abaixo de 1 é metodologicamente intencional: a seleção obedece ao ranking das regras elegíveis, sem ocultar associações negativas ou próximas da independência."),
        ("33 Regras óbvias", f"Foram classificadas {classes['ÓBVIA']} regras como óbvias. Incluem direções inversas do mesmo itemset, associações quase independentes e casos explicáveis pela frequência dominante de casado/união civil."),
        ("34 Regras interessantes", f"Foram classificadas {classes['INTERESSANTE']} regras como interessantes. Elas conectam dimensões distintas e apresentam Lift positivo pequeno, mas exigem cautela pela baixa cobertura da amostra."),
        ("35 Novidades identificadas", f"Foi identificada {classes['NOVIDADE']} novidade preliminar: crédito muito alto e situação familiar separada/viúva/não informada associados a um a dois créditos ativos. O padrão tem Lift 1,10, confiança 48% e suporte 1%; é hipótese para investigação, não conclusão definitiva."),
        ("36 Interpretação comercial", "As regras podem orientar hipóteses de segmentação e perguntas para análise posterior, como investigar perfis de crédito ativo por situação familiar. Elas não devem ser usadas isoladamente para aprovar, negar ou precificar crédito, pois são associações de uma amostra e não modelos causais."),
        ("37 Visualizações", "Foram produzidos gráficos de Lift, confiança versus Lift, distribuição das classificações e frequência dos itens. O gráfico de suporte versus confiança não foi usado porque o suporte é constante em 1% nas vinte regras, o que não adicionaria informação."),
        ("38 Limitações", "As principais limitações são a amostra de 10.000 linhas, discretização que reduz granularidade, suporte baixo das regras selecionadas, Lift próximo de 1 em grande parte do ranking e a limitação do WEKA de não filtrar regras por total de itens."),
        ("39 Ética e uso responsável", "Características financeiras e familiares exigem cuidado. As associações não devem sustentar discriminação, decisões automatizadas sem validação, inferências sobre indivíduos ou tratamento desigual. Qualquer uso operacional exige avaliação de vieses, governança, explicabilidade e supervisão humana."),
        ("40 Conclusão", "O trabalho cumpriu a cadeia auditável de discretização, codificação binária, mineração real, itemsets, fechamento, seleção e interpretação. O resultado principal é metodológico: o conjunto oferece hipóteses de associação, mas a força limitada e o suporte de 1% recomendam validação adicional antes de qualquer uso prático."),
    ]
    selected = [[r['Posição'], r['Antecedente'], r['Consequente'], r['Suporte'], r['Confiança'], r['Lift'], r['Classificação']] for r in top]
    support_rows = [[r['teste'], r['suporte_efetivo_reportado'], r['num_rules_configurado'], r['metrica'], r['regras_geradas'], r['regras_tres_itens'], r['itemsets']] for r in support]
    disc_rows = [[r.get('atributo',''), r.get('categoria',''), r.get('limite_inferior',''), r.get('limite_superior','')] for r in decisions]
    autores = "Gil Antony Borba; Raphael Muniz Varela; Victor Leal; Ygor Marangoni"
    md = ["# UNIVERSIDADE FEDERAL DE UBERLÂNDIA", "", "## FACULDADE DE COMPUTAÇÃO", "", "# TRABALHO PRÁTICO DE CIÊNCIA DE DADOS", "", "## REGRAS DE ASSOCIAÇÃO APRIORI", "", autores + "  ", "Monte Carmelo - MG, 2026", "", "---", "", "# FOLHA DE ROSTO", "", "Trabalho Prático de Ciência de Dados - Regras de Associação Apriori.", "", autores + ". Trabalho apresentado à disciplina de Ciência de Dados da Universidade Federal de Uberlândia, como requisito parcial de avaliação. Professor: Carlos Cesar Mansur Tuma.", "", "---", "", "# SUMÁRIO", "", "O sumário automático será atualizado no Microsoft Word durante a preparação final do arquivo.", ""]
    for title, paragraph in section_text:
        md += [f"# {title}", "", paragraph, ""]
        if title.startswith("7 "):
            md += [table_md(["Dimensão", "Representação"], [["FAIXA_CREDITO","quatro faixas"],["FAIXA_RENDA","quatro faixas"],["FAIXA_IDADE","cinco faixas"],["CATEGORIA_FILHOS","três categorias"],["POSSE_CARRO","com/sem carro"],["SITUACAO_FAMILIAR","três grupos"],["FAIXA_CREDITOS_ATIVOS","três faixas"],["FAIXA_TAXA_REJEICAO","três faixas"]]), ""]
        if title.startswith("11 "):
            md += [table_md(["Atributo", "Categoria", "Limite inferior", "Limite superior"], disc_rows), ""]
        if title.startswith("23 "):
            md += [table_md(["Teste", "Suporte", "N", "Métrica", "Regras", "3 itens", "Itemsets"], support_rows), ""]
        if title.startswith("32 "):
            md += [table_md(["Pos.", "Antecedente", "Consequente", "Suporte", "Confiança", "Lift", "Classe"], selected), ""]
        if title.startswith("37 "):
            md += ["![Figura 1 - Top 20 por Lift](imagens/regras/01_top20_lift.png)", "", "![Figura 2 - Confiança e Lift](imagens/regras/02_confianca_lift_classificacao.png)", "", "![Figura 3 - Distribuição das classificações](imagens/regras/03_distribuicao_classificacao.png)", "", "![Figura 4 - Frequência dos itens](imagens/regras/04_frequencia_itens_top20.png)", ""]
    md += ["# 41 Referências", "", "AGRAWAL, R.; SRIKANT, R. Fast algorithms for mining association rules. Proceedings of the 20th VLDB Conference, 1994.", "", "HAN, J.; KAMBER, M.; PEI, J. Data Mining: Concepts and Techniques. 4. ed. Morgan Kaufmann, 2023.", "", "WEKA. Waikato Environment for Knowledge Analysis, versão 3.8.7. Documentação e ajuda da instalação local.", "", "# 42 Apêndices", "", "Apêndice A - Saídas integrais do WEKA, configurações, testes de suporte, itemsets, conjunto fechado e CSVs de regras permanecem preservados no repositório."]
    md_path = REL / "relatorio_regras_associacao.md"; md_path.write_text("\n".join(md), encoding="utf-8")

    doc = Document(); sec = doc.sections[0]
    for s in doc.sections:
        s.page_width = Mm(210); s.page_height = Mm(297)
        s.top_margin = Cm(3); s.left_margin = Cm(3); s.right_margin = Cm(2); s.bottom_margin = Cm(2)
    normal = doc.styles['Normal']; normal.font.name = 'Times New Roman'; normal.font.size = Pt(12); normal.paragraph_format.space_after = Pt(0); normal.paragraph_format.line_spacing = 1.5; normal.paragraph_format.first_line_indent = Cm(1.25); normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for style, size in [('Heading 1',16),('Heading 2',13),('Heading 3',12)]:
        st=doc.styles[style]; st.font.name='Times New Roman'; st.font.size=Pt(12); st.font.bold=True; st.paragraph_format.space_before=Pt(12); st.paragraph_format.space_after=Pt(6); st.paragraph_format.line_spacing=1.5
    settings = doc.settings.element
    update = OxmlElement('w:updateFields'); update.set(qn('w:val'), 'true'); settings.append(update)
    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT; add_field(footer, 'PAGE')
    cover = doc.add_paragraph(); cover.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for text,size,bold in [("UNIVERSIDADE FEDERAL DE UBERLÂNDIA",12,False),("FACULDADE DE COMPUTAÇÃO",12,False),("",12,False),("TRABALHO PRÁTICO DE CIÊNCIA DE DADOS",16,True),("REGRAS DE ASSOCIAÇÃO APRIORI",18,True),("",12,False),("Gil Antony Borba",12,False),("Raphael Muniz Varela",12,False),("Victor Leal",12,False),("Ygor Marangoni",12,False),("",12,False),("Monte Carmelo - MG, 2026",12,False)]:
        r=cover.add_run(text+'\n'); r.font.name='Times New Roman'; r.font.size=Pt(12); r.bold=bold
    doc.add_page_break(); doc.add_heading('FOLHA DE ROSTO', 1); doc.add_paragraph('Gil Antony Borba, Raphael Muniz Varela, Victor Leal e Ygor Marangoni. Trabalho Prático de Ciência de Dados - Regras de Associação Apriori. Trabalho apresentado à disciplina de Ciência de Dados da Universidade Federal de Uberlândia, como requisito parcial de avaliação. Professor: Carlos Cesar Mansur Tuma.')
    doc.add_page_break(); doc.add_heading('SUMÁRIO', 1); toc=doc.add_paragraph(); add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    for title, paragraph in section_text:
        doc.add_heading(title, 1); doc.add_paragraph(paragraph)
        if title.startswith('7 '): add_doc_table(doc,['Dimensão','Representação'],[["FAIXA_CREDITO","quatro faixas"],["FAIXA_RENDA","quatro faixas"],["FAIXA_IDADE","cinco faixas"],["CATEGORIA_FILHOS","três categorias"],["POSSE_CARRO","com/sem carro"],["SITUACAO_FAMILIAR","três grupos"],["FAIXA_CREDITOS_ATIVOS","três faixas"],["FAIXA_TAXA_REJEICAO","três faixas"]])
        if title.startswith('11 '): add_doc_table(doc,['Atributo','Categoria','Inf.','Sup.'],disc_rows,7)
        if title.startswith('23 '): add_doc_table(doc,['Teste','Sup.','N','Métrica','Regras','3 itens','Itemsets'],support_rows,6)
        if title.startswith('32 '): add_doc_table(doc,['Pos.','Antecedente','Consequente','Sup.','Conf.','Lift','Classe'],selected,6)
        if title.startswith('37 '):
            for i,name in enumerate(['01_top20_lift.png','02_confianca_lift_classificacao.png','03_distribuicao_classificacao.png','04_frequencia_itens_top20.png'],1):
                p=doc.add_paragraph(f'Figura {i} - Visualização dos resultados.'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Cm(0)
                doc.add_picture(str(REL/'imagens/regras'/name), width=Cm(15.5)); src=doc.add_paragraph('Fonte: elaboração própria a partir da saída do WEKA 3.8.7.'); src.paragraph_format.first_line_indent=Cm(0); src.runs[0].italic=True
    doc.add_heading('41 Referências',1)
    for ref in ["AGRAWAL, R.; SRIKANT, R. Fast algorithms for mining association rules. Proceedings of the 20th VLDB Conference, 1994.","HAN, J.; KAMBER, M.; PEI, J. Data Mining: Concepts and Techniques. 4. ed. Morgan Kaufmann, 2023.","WEKA. Waikato Environment for Knowledge Analysis, versão 3.8.7. Documentação e ajuda da instalação local."]: doc.add_paragraph(ref)
    doc.add_heading('42 Apêndices',1); doc.add_paragraph('Apêndice A - Arquivos auditáveis preservados no repositório: saídas do WEKA, configurações, testes de suporte, itemsets, conjunto fechado e CSVs de regras.')
    doc.save(REL / 'relatorio_regras_associacao_abnt.docx')
    print(f'MD={md_path}')
    print(f'DOCX={REL / "relatorio_regras_associacao_abnt.docx"}')
    print(f'TOP20={len(top)}; TESTES={len(support)}; ITEMSETS_FECHADOS={n_closed}')

if __name__ == '__main__': main()
